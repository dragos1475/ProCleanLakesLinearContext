from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
import os
from datetime import datetime
from jinja2 import Environment, FileSystemLoader
from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error, explained_variance_score
from docx import Document
from docx.shared import Inches
import io
from jinja2 import Template
import glob
import shutil

class ReportGenerator:
    def __init__(self, filename="regression_report.pdf", output_dir='reports'):
        self.filename = filename
        self.doc = SimpleDocTemplate(filename, pagesize=letter)
        self.story = []
        self.styles = getSampleStyleSheet()
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceAfter=30
        )
        self.subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceAfter=20
        )
        self.output_dir = output_dir
        self.env = Environment(loader=FileSystemLoader('templates'))
        os.makedirs(output_dir, exist_ok=True)
    
    def add_title(self, text):
        self.story.append(Paragraph(text, self.title_style))
        self.story.append(Spacer(1, 12))
    
    def add_subtitle(self, text):
        self.story.append(Paragraph(text, self.subtitle_style))
        self.story.append(Spacer(1, 12))
    
    def add_text(self, text):
        self.story.append(Paragraph(text, self.styles["Normal"]))
        self.story.append(Spacer(1, 12))
    
    def add_correlation_matrix(self, corr_matrix):
        self.add_subtitle("Correlation Matrix")
        
        # Convert correlation matrix to DataFrame for better formatting
        df = pd.DataFrame(corr_matrix).round(2)
        data = [df.columns.tolist()] + df.values.tolist()
        
        # Create table
        t = Table(data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        self.story.append(t)
        self.story.append(Spacer(1, 12))
    
    def add_pca_analysis(self, pca, feature_names):
        self.add_subtitle("PCA Analysis")
        
        # Explained variance ratio
        self.add_text("Explained Variance Ratio:")
        data = [['Component', 'Explained Variance', 'Cumulative Variance']]
        cumulative = 0
        for i, var in enumerate(pca.explained_variance_ratio_):
            cumulative += var
            data.append([f'PC{i+1}', f'{var:.2f}', f'{cumulative:.2f}'])
        
        t = Table(data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        self.story.append(t)
        self.story.append(Spacer(1, 12))
        
        # Component loadings
        self.add_text("Component Loadings:")
        loadings = pd.DataFrame(
            pca.components_,
            columns=feature_names,
            index=[f'PC{i+1}' for i in range(len(pca.components_))]
        ).round(2)
        data = [[''] + loadings.columns.tolist()] + \
               [[index] + row.tolist() for index, row in loadings.iterrows()]
        
        t = Table(data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        self.story.append(t)
        self.story.append(Spacer(1, 12))
    
    def add_model_equation(self, model_name, coefficients, intercept, feature_names):
        self.add_subtitle(f"{model_name} Model Equation")
        
        equation = f"y = {intercept:.2f}"
        for i, coef in enumerate(coefficients):
            equation += f" + {coef:.2f} * {feature_names[i]}"
        
        self.add_text(equation)
        self.story.append(Spacer(1, 12))
    
    def add_model_metrics(self, model_name, metrics):
        self.add_subtitle(f"{model_name} Performance Metrics")
        
        data = [['Metric', 'Value']]
        for metric, value in metrics.items():
            data.append([metric, f'{value:.2f}'])
        
        t = Table(data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        self.story.append(t)
        self.story.append(Spacer(1, 12))
    
    def add_plot(self, plot_func, *args, **kwargs):
        """Add a plot to the report"""
        # Create a BytesIO object to save the plot
        buf = BytesIO()
        plot_func(*args, **kwargs)
        plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        
        # Add the image to the report
        img = Image(buf, width=6*inch, height=4*inch)
        self.story.append(img)
        self.story.append(Spacer(1, 12))
    
    def generate_report(self, model_name, metrics, feature_importance, correlation_matrix, 
                       pca_results, residual_analysis, actual_vs_predicted, residuals_plot, 
                       error_distribution, pca_variance_plot, correlation_heatmap):
        """Generate HTML report with all analysis results"""
        # Create reports directory if it doesn't exist
        os.makedirs('reports', exist_ok=True)
        
        # Create HTML content using Jinja2 template
        template = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>{{ model_name }} Analysis Report</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                h1, h2 { color: #2c3e50; }
                .section { margin-bottom: 30px; }
                .metric { margin: 10px 0; }
                .plot { margin: 20px 0; }
                table { border-collapse: collapse; width: 100%; margin: 20px 0; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f2f2f2; }
                .good { color: green; }
                .warning { color: orange; }
                .error { color: red; }
            </style>
        </head>
        <body>
            <h1>{{ model_name }} Analysis Report</h1>
            
            <div class="section">
                <h2>Model Metrics</h2>
                <table>
                    <tr>
                        <th>Metric</th>
                        <th>Value</th>
                        <th>Interpretation</th>
                    </tr>
                    <tr>
                        <td>R² Score</td>
                        <td>{{ metrics.R2 }}</td>
                        <td>{% if metrics.R2 >= 0.7 %}<span class="good">Good fit</span>
                            {% elif metrics.R2 >= 0.5 %}<span class="warning">Moderate fit</span>
                            {% else %}<span class="error">Poor fit</span>{% endif %}</td>
                    </tr>
                    <tr>
                        <td>RMSE</td>
                        <td>{{ metrics.RMSE }}</td>
                        <td>{% if metrics.RMSE <= 0.1 %}<span class="good">Low error</span>
                            {% elif metrics.RMSE <= 0.3 %}<span class="warning">Moderate error</span>
                            {% else %}<span class="error">High error</span>{% endif %}</td>
                    </tr>
                    <tr>
                        <td>MAE</td>
                        <td>{{ metrics.MAE }}</td>
                        <td>{% if metrics.MAE <= 0.1 %}<span class="good">Low error</span>
                            {% elif metrics.MAE <= 0.3 %}<span class="warning">Moderate error</span>
                            {% else %}<span class="error">High error</span>{% endif %}</td>
                    </tr>
                    <tr>
                        <td>Explained Variance</td>
                        <td>{{ metrics['Explained Variance'] }}</td>
                        <td>{% if metrics['Explained Variance'] >= 0.7 %}<span class="good">Good explanation</span>
                            {% elif metrics['Explained Variance'] >= 0.5 %}<span class="warning">Moderate explanation</span>
                            {% else %}<span class="error">Poor explanation</span>{% endif %}</td>
                    </tr>
                </table>
            </div>
            
            <div class="section">
                <h2>Feature Importance</h2>
                <table>
                    <tr>
                        <th>Feature</th>
                        <th>Importance</th>
                    </tr>
                    {% for feature, importance in feature_importance.items() %}
                    <tr>
                        <td>{{ feature }}</td>
                        <td>{{ "%.4f"|format(importance) }}</td>
                    </tr>
                    {% endfor %}
                </table>
            </div>
            
            <div class="section">
                <h2>Residual Analysis</h2>
                <table>
                    <tr>
                        <th>Test</th>
                        <th>Statistic</th>
                        <th>p-value</th>
                        <th>Interpretation</th>
                    </tr>
                    <tr>
                        <td>Shapiro-Wilk (Normality)</td>
                        <td>{{ residual_analysis.normality.shapiro_wilk.statistic }}</td>
                        <td>{{ residual_analysis.normality.shapiro_wilk.p_value }}</td>
                        <td>{% if residual_analysis.normality.shapiro_wilk.p_value > 0.05 %}
                            <span class="good">Normal</span>
                            {% else %}<span class="error">Not normal</span>{% endif %}</td>
                    </tr>
                    <tr>
                        <td>Breusch-Pagan (Heteroscedasticity)</td>
                        <td>-</td>
                        <td>{{ residual_analysis.heteroscedasticity.breusch_pagan_p_value }}</td>
                        <td>{% if residual_analysis.heteroscedasticity.breusch_pagan_p_value > 0.05 %}
                            <span class="good">Homoscedastic</span>
                            {% else %}<span class="error">Heteroscedastic</span>{% endif %}</td>
                    </tr>
                    <tr>
                        <td>Durbin-Watson (Autocorrelation)</td>
                        <td>{{ residual_analysis.autocorrelation.durbin_watson }}</td>
                        <td>-</td>
                        <td>{% if 1.5 <= residual_analysis.autocorrelation.durbin_watson <= 2.5 %}
                            <span class="good">No autocorrelation</span>
                            {% else %}<span class="error">Autocorrelation present</span>{% endif %}</td>
                    </tr>
                </table>
            </div>
            
            <div class="section">
                <h2>Visualizations</h2>
                <div class="plot">
                    <h3>Actual vs Predicted Values</h3>
                    <img src="{{ actual_vs_predicted }}" alt="Actual vs Predicted Values">
                </div>
                <div class="plot">
                    <h3>Residuals Plot</h3>
                    <img src="{{ residuals_plot }}" alt="Residuals Plot">
                </div>
                <div class="plot">
                    <h3>Error Distribution</h3>
                    <img src="{{ error_distribution }}" alt="Error Distribution">
                </div>
                <div class="plot">
                    <h3>PCA Explained Variance</h3>
                    <img src="{{ pca_variance_plot }}" alt="PCA Explained Variance">
                </div>
                <div class="plot">
                    <h3>Feature Correlation Heatmap</h3>
                    <img src="{{ correlation_heatmap }}" alt="Feature Correlation Heatmap">
                </div>
            </div>
        </body>
        </html>
        """
        
        # Render template with data
        html_content = Template(template).render(
            model_name=model_name,
            metrics=metrics,
            feature_importance=feature_importance,
            residual_analysis=residual_analysis,
            actual_vs_predicted=actual_vs_predicted,
            residuals_plot=residuals_plot,
            error_distribution=error_distribution,
            pca_variance_plot=pca_variance_plot,
            correlation_heatmap=correlation_heatmap
        )
        
        # Save HTML report
        output_path = f'reports/{model_name}_report.html'
        with open(output_path, 'w') as f:
            f.write(html_content)
        
        return output_path
    
    def generate_summary(self, model_name, metrics, feature_importance, residual_analysis):
        """Print a console summary of the model, metrics, feature importance, and residual analysis."""
        print(f"\n=== {model_name} Summary ===")
        print("\nMetrics:")
        for metric, value in metrics.items():
            if isinstance(value, (int, float)):
                print(f"{metric}: {value:.4f}")
            else:
                print(f"{metric}: {value}")
        
        print("\nFeature Importance:")
        for feature, importance in feature_importance.items():
            print(f"{feature}: {importance:.4f}")
        
        print("\nResidual Analysis:")
        for test, result in residual_analysis.items():
            if isinstance(result, dict):
                print(f"\n{test}:")
                for sub_test, sub_result in result.items():
                    if isinstance(sub_result, dict):
                        print(f"  {sub_test}:")
                        for stat, value in sub_result.items():
                            if isinstance(value, (int, float)):
                                print(f"    {stat}: {value:.4f}")
                            else:
                                print(f"    {stat}: {value}")
                    else:
                        if isinstance(sub_result, (int, float)):
                            print(f"  {sub_test}: {sub_result:.4f}")
                        else:
                            print(f"  {sub_test}: {sub_result}")
            else:
                if isinstance(result, (int, float)):
                    print(f"{test}: {result:.4f}")
                else:
                    print(f"{test}: {result}")

    def generate_pdf_report(self):
        """Generate and save the PDF report"""
        self.doc.build(self.story)

    def cleanup_temp_files(self):
        """Remove temporary files after report generation, preserving plot files needed for reports."""
        # Only remove files that are not needed for reports
        patterns_to_clean = [
            '*.joblib',  # Model files
            '*.log',     # Log files
            '*.tmp',     # Temporary files
            '*.pyc',     # Python cache files
            '*.DS_Store'  # macOS system files
        ]
        
        # Clean only non-report files in the main and reports directories
        directories = ['.', self.output_dir]
        
        for directory in directories:
            # Clean files
            for pattern in patterns_to_clean:
                for file in glob.glob(os.path.join(directory, pattern)):
                    try:
                        os.remove(file)
                        print(f"Removed temporary file: {file}")
                    except Exception as e:
                        print(f"Warning: Could not remove {file}: {e}")
            
            # Handle __pycache__ directory separately
            pycache_dir = os.path.join(directory, '__pycache__')
            if os.path.exists(pycache_dir):
                try:
                    shutil.rmtree(pycache_dir)
                    print(f"Removed directory: {pycache_dir}")
                except Exception as e:
                    print(f"Warning: Could not remove {pycache_dir}: {e}")

    def generate_combined_word_report(self, raw_results, standardized_results, target_name):
        """Generate a Word document containing reports for both raw and standardized data analysis."""
        doc = Document()
        
        # Add title
        doc.add_heading('Regression Analysis Report - Raw and Standardized Data', 0)
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Add correlation analysis section (using raw data only)
        first_model_raw = next(iter(raw_results.values()))
        
        doc.add_heading('Feature Correlation Analysis', level=1)
        
        # Raw data correlation (including target variable)
        doc.add_heading('Correlation Matrix', level=2)
        correlation_matrix = first_model_raw['correlation_matrix']
        correlation_table = doc.add_table(rows=1, cols=len(correlation_matrix.columns) + 1)
        correlation_table.style = 'Table Grid'
        
        # Add headers
        header_cells = correlation_table.rows[0].cells
        header_cells[0].text = 'Feature'
        for i, col in enumerate(correlation_matrix.columns):
            header_cells[i + 1].text = col
        
        # Add correlation values
        for i, (index, row) in enumerate(correlation_matrix.iterrows()):
            row_cells = correlation_table.add_row().cells
            row_cells[0].text = index
            for j, value in enumerate(row):
                row_cells[j + 1].text = f'{value:.3f}'
        
        # Add correlation heatmap
        if 'correlation_heatmap' in first_model_raw['plots']:
            doc.add_paragraph('\nCorrelation Heatmap:')
            doc.add_picture(first_model_raw['plots']['correlation_heatmap'], width=Inches(6))
        
        # Add PCA analysis section
        doc.add_heading('Principal Component Analysis', level=1)
        
        # Raw data PCA
        pca_results = first_model_raw['pca_results']
        
        # Add explained variance ratio
        doc.add_heading('Explained Variance Ratio', level=2)
        pca_table = doc.add_table(rows=1, cols=3)
        pca_table.style = 'Table Grid'
        header_cells = pca_table.rows[0].cells
        header_cells[0].text = 'Component'
        header_cells[1].text = 'Explained Variance'
        header_cells[2].text = 'Cumulative Variance'
        
        cumulative_variance = 0
        for i, var in enumerate(pca_results['explained_variance_ratio']):
            cumulative_variance += var
            row_cells = pca_table.add_row().cells
            row_cells[0].text = f'PC{i+1}'
            row_cells[1].text = f'{var:.4f}'
            row_cells[2].text = f'{cumulative_variance:.4f}'
        
        # Add PCA variance plot
        if 'pca_variance' in first_model_raw['plots']:
            doc.add_paragraph('\nPCA Explained Variance:')
            doc.add_picture(first_model_raw['plots']['pca_variance'], width=Inches(6))
        
        # Add component loadings
        doc.add_heading('Component Loadings', level=2)
        loadings_table = doc.add_table(rows=1, cols=len(pca_results['components']) + 1)
        loadings_table.style = 'Table Grid'
        
        # Add headers
        header_cells = loadings_table.rows[0].cells
        header_cells[0].text = 'Feature'
        for i in range(len(pca_results['components'])):
            header_cells[i + 1].text = f'PC{i+1}'
        
        # Add loadings values
        features = [col for col in correlation_matrix.columns if col != target_name]
        for i, feature in enumerate(features):
            row_cells = loadings_table.add_row().cells
            row_cells[0].text = feature
            for j in range(len(pca_results['components'])):
                row_cells[j + 1].text = f'{pca_results["components"][j][i]:.4f}'
        
        # Add model results sections
        for data_type, results in [("Raw Data", raw_results), ("Standardized Data", standardized_results)]:
            doc.add_heading(f'\nModel Results for {data_type}', level=1)
            
            for model_name, model_results in results.items():
                doc.add_heading(f'\n{model_name}', level=2)
                
                # Model metrics
                doc.add_heading('Model Metrics', level=3)
                metrics_table = doc.add_table(rows=1, cols=2)
                metrics_table.style = 'Table Grid'
                metrics_table.rows[0].cells[0].text = 'Metric'
                metrics_table.rows[0].cells[1].text = 'Value'
                
                for metric, value in model_results['metrics'].items():
                    if metric not in ['coefficients', 'intercept', 'scaling_params']:
                        row = metrics_table.add_row()
                        row.cells[0].text = metric
                        row.cells[1].text = f'{value:.4f}' if isinstance(value, (int, float)) else str(value)
                
                # Model equation
                doc.add_heading('Model Equation', level=3)
                coefficients = model_results['metrics']['coefficients']
                intercept = model_results['metrics']['intercept']
                
                if data_type == "Raw Data":
                    equation = f"{target_name} = {intercept:.4f}"
                    for feature, coef in coefficients.items():
                        equation += f" + ({coef:.4f} × {feature})"
                else:
                    equation = f"Z_{target_name} = {intercept:.4f}"
                    for feature, coef in coefficients.items():
                        equation += f" + ({coef:.4f} × Z_{feature})"
                
                doc.add_paragraph(equation)
                
                if data_type == "Standardized Data":
                    # Add standardization parameters
                    doc.add_heading('Standardization Parameters', level=3)
                    scaling_params = model_results['metrics']['scaling_params']
                    
                    doc.add_paragraph("Feature Means:")
                    for feature, mean in scaling_params['feature_means'].items():
                        doc.add_paragraph(f"• {feature}: {mean:.4f}")
                    
                    doc.add_paragraph("\nFeature Standard Deviations:")
                    for feature, std in scaling_params['feature_stds'].items():
                        doc.add_paragraph(f"• {feature}: {std:.4f}")
                    
                    doc.add_paragraph(f"\nTarget Mean: {scaling_params['target_mean']:.4f}")
                    doc.add_paragraph(f"Target Standard Deviation: {scaling_params['target_std']:.4f}")
                
                # Feature importance
                doc.add_heading('Feature Importance', level=3)
                importance_table = doc.add_table(rows=1, cols=2)
                importance_table.style = 'Table Grid'
                importance_table.rows[0].cells[0].text = 'Feature'
                importance_table.rows[0].cells[1].text = 'Importance'
                
                for feature, importance in model_results['feature_importance'].items():
                    row = importance_table.add_row()
                    row.cells[0].text = feature
                    row.cells[1].text = f'{importance:.4f}'
                
                # Residual analysis
                doc.add_heading('Residual Analysis', level=3)
                residual_table = doc.add_table(rows=1, cols=3)
                residual_table.style = 'Table Grid'
                header_cells = residual_table.rows[0].cells
                header_cells[0].text = 'Test'
                header_cells[1].text = 'Statistic'
                header_cells[2].text = 'p-value'
                
                for test_name, test_results in model_results['residual_analysis'].items():
                    if isinstance(test_results, dict):
                        for sub_test, sub_results in test_results.items():
                            row = residual_table.add_row()
                            row.cells[0].text = f'{test_name} - {sub_test}'
                            if isinstance(sub_results, dict):
                                row.cells[1].text = f"{sub_results.get('statistic', 'N/A')}"
                                row.cells[2].text = f"{sub_results.get('p_value', 'N/A')}"
                            else:
                                row.cells[1].text = f"{sub_results}"
                                row.cells[2].text = "N/A"
                
                # Add model-specific plots (excluding correlation and PCA plots)
                doc.add_heading('Model Visualizations', level=3)
                for plot_name, plot_path in model_results['plots'].items():
                    if plot_name not in ['correlation_heatmap', 'pca_variance'] and plot_path and os.path.exists(plot_path):
                        doc.add_paragraph(f'\n{plot_name.replace("_", " ").title()}:')
                        doc.add_picture(plot_path, width=Inches(6))
        
        # Save the document
        output_path = os.path.join(self.output_dir, 'combined_regression_analysis.docx')
        doc.save(output_path)
        
        # Clean up temporary files
        self.cleanup_temp_files()
        
        return output_path

    def generate_html_report(self, model_name, data_type, results, target_name):
        """Generate HTML report for a specific model and data type (raw or standardized)"""
        # Create reports directory if it doesn't exist
        os.makedirs('reports', exist_ok=True)
        
        # Get target_std from metrics if available (standardized data) or calculate it for raw data
        if 'scaling_params' in results['metrics']:
            target_std = results['metrics']['scaling_params']['target_std']
        else:
            # For raw data, we'll use more generous thresholds
            target_std = 1.0  # Using 1.0 as a baseline for raw data
        
        # Add target_std to metrics for template use
        results['metrics']['target_std'] = target_std
        
        # Update plot paths to use the correct path (plots directory)
        plots = {}
        for plot_name, plot_path in results['plots'].items():
            if plot_path:
                plots[plot_name] = os.path.join('plots', os.path.basename(plot_path))
        
        # Prepare template data
        template_data = {
            'model_name': model_name,
            'data_type': data_type,
            'target_name': target_name,
            'metrics': results['metrics'],
            'model_equation': results['model_equation'],
            'feature_importance': results['feature_importance'],
            'plots': plots
        }
        
        # Load and customize the template
        template_str = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{{ model_name }} - {{ data_type }} Analysis Report</title>
            <style>
                body { 
                    font-family: Arial, sans-serif; 
                    margin: 20px; 
                    line-height: 1.6;
                    color: #333;
                    max-width: 1200px;
                    margin: 0 auto;
                    padding: 20px;
                }
                h1, h2 { 
                    color: #2c3e50; 
                    border-bottom: 2px solid #eee;
                    padding-bottom: 10px;
                }
                .section { 
                    margin-bottom: 40px; 
                    background: #fff;
                    padding: 20px;
                    border-radius: 8px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                }
                table { 
                    border-collapse: collapse; 
                    width: 100%; 
                    margin: 20px 0;
                    background: #fff;
                }
                th, td { 
                    border: 1px solid #ddd; 
                    padding: 12px; 
                    text-align: left; 
                }
                th { 
                    background-color: #f8f9fa;
                    font-weight: bold;
                }
                tr:nth-child(even) { background-color: #f8f9fa; }
                .good { color: #28a745; font-weight: bold; }
                .warning { color: #ffc107; font-weight: bold; }
                .error { color: #dc3545; font-weight: bold; }
                .plot { 
                    margin: 20px 0;
                    text-align: center;
                }
                .plot img {
                    max-width: 100%;
                    height: auto;
                    border-radius: 4px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                }
                .equation {
                    background: #f8f9fa;
                    padding: 15px;
                    border-radius: 4px;
                    font-family: monospace;
                    margin: 20px 0;
                }
                .feature-importance {
                    display: flex;
                    flex-direction: column;
                    gap: 10px;
                }
                .feature-bar {
                    display: flex;
                    align-items: center;
                    gap: 10px;
                }
                .bar {
                    background: #007bff;
                    height: 20px;
                    border-radius: 4px;
                }
            </style>
        </head>
        <body>
            <h1>{{ model_name }} - {{ data_type }} Analysis Report</h1>
            
            <div class="section">
                <h2>Model Metrics</h2>
                <table>
                    <tr>
                        <th>Metric</th>
                        <th>Value</th>
                        <th>Interpretation</th>
                    </tr>
                    <tr>
                        <td>R² Score</td>
                        <td>{{ "%.4f"|format(metrics.R2) }}</td>
                        <td>{% if metrics.R2 >= 0.7 %}<span class="good">Good fit</span>
                            {% elif metrics.R2 >= 0.5 %}<span class="warning">Moderate fit</span>
                            {% else %}<span class="error">Poor fit</span>{% endif %}</td>
                    </tr>
                    <tr>
                        <td>RMSE</td>
                        <td>{{ "%.4f"|format(metrics.RMSE) }}</td>
                        <td>{% if data_type == 'Standardized' %}
                            {% if metrics.RMSE <= 0.3 %}<span class="good">Low error</span>
                            {% elif metrics.RMSE <= 0.5 %}<span class="warning">Moderate error</span>
                            {% else %}<span class="error">High error</span>{% endif %}
                            {% else %}
                            {% if metrics.RMSE <= metrics.target_std * 0.3 %}<span class="good">Low error</span>
                            {% elif metrics.RMSE <= metrics.target_std * 0.5 %}<span class="warning">Moderate error</span>
                            {% else %}<span class="error">High error</span>{% endif %}
                            {% endif %}</td>
                    </tr>
                    <tr>
                        <td>MAE</td>
                        <td>{{ "%.4f"|format(metrics.MAE) }}</td>
                        <td>{% if data_type == 'Standardized' %}
                            {% if metrics.MAE <= 0.25 %}<span class="good">Low error</span>
                            {% elif metrics.MAE <= 0.4 %}<span class="warning">Moderate error</span>
                            {% else %}<span class="error">High error</span>{% endif %}
                            {% else %}
                            {% if metrics.MAE <= metrics.target_std * 0.25 %}<span class="good">Low error</span>
                            {% elif metrics.MAE <= metrics.target_std * 0.4 %}<span class="warning">Moderate error</span>
                            {% else %}<span class="error">High error</span>{% endif %}
                            {% endif %}</td>
                    </tr>
                    <tr>
                        <td>Explained Variance</td>
                        <td>{{ "%.4f"|format(metrics['Explained Variance']) }}</td>
                        <td>{% if metrics['Explained Variance'] >= 0.7 %}<span class="good">Good explanation</span>
                            {% elif metrics['Explained Variance'] >= 0.5 %}<span class="warning">Moderate explanation</span>
                            {% else %}<span class="error">Poor explanation</span>{% endif %}</td>
                    </tr>
                </table>
                
                {% if data_type == 'Standardized' and metrics.scaling_params %}
                <h3>Standardization Parameters</h3>
                <table>
                    <tr>
                        <th>Parameter</th>
                        <th>Value</th>
                    </tr>
                    <tr>
                        <td>Target Mean</td>
                        <td>{{ "%.4f"|format(metrics.scaling_params.target_mean) }}</td>
                    </tr>
                    <tr>
                        <td>Target Standard Deviation</td>
                        <td>{{ "%.4f"|format(metrics.scaling_params.target_std) }}</td>
                    </tr>
                </table>
                
                <h4>Feature Scaling Parameters</h4>
                <table>
                    <tr>
                        <th>Feature</th>
                        <th>Mean</th>
                        <th>Standard Deviation</th>
                    </tr>
                    {% for feature in metrics.scaling_params.feature_means.keys() %}
                    <tr>
                        <td>{{ feature }}</td>
                        <td>{{ "%.4f"|format(metrics.scaling_params.feature_means[feature]) }}</td>
                        <td>{{ "%.4f"|format(metrics.scaling_params.feature_stds[feature]) }}</td>
                    </tr>
                    {% endfor %}
                </table>
                {% endif %}
            </div>

            <div class="section">
                <h2>Model Equation</h2>
                <div class="equation">
                    {% if data_type == 'Raw' %}
                    {{ target_name }} = {{ "%.4f"|format(model_equation.intercept) }}
                    {% for feature, coef in model_equation.coefficients.items() %}
                    {% if coef >= 0 %} + {% endif %}{{ "%.4f"|format(coef) }} &times; {{ feature }}
                    {% endfor %}
                    {% else %}
                    Z_{{ target_name }} = {{ "%.4f"|format(model_equation.intercept) }}
                    {% for feature, coef in model_equation.coefficients.items() %}
                    {% if coef >= 0 %} + {% endif %}{{ "%.4f"|format(coef) }} &times; Z_{{ feature }}
                    {% endfor %}
                    {% endif %}
                </div>
            </div>

            <div class="section">
                <h2>Feature Importance</h2>
                <div class="feature-importance">
                    {% for feature, importance in feature_importance.items() %}
                    <div class="feature-bar">
                        <span style="width: 200px;">{{ feature }}</span>
                        <div class="bar" style="width: {{ importance * 100 }}%;"></div>
                        <span>{{ "%.4f"|format(importance) }}</span>
                    </div>
                    {% endfor %}
                </div>
            </div>

            <div class="section">
                <h2>Model Performance Visualizations</h2>
                {% if plots.actual_vs_predicted %}
                <div class="plot">
                    <h3>Actual vs Predicted Values</h3>
                    <img src="{{ plots.actual_vs_predicted }}" alt="Actual vs Predicted Values">
                </div>
                {% endif %}
                {% if plots.residuals %}
                <div class="plot">
                    <h3>Residuals Plot</h3>
                    <img src="{{ plots.residuals }}" alt="Residuals Plot">
                </div>
                {% endif %}
                {% if plots.error_distribution %}
                <div class="plot">
                    <h3>Error Distribution</h3>
                    <img src="{{ plots.error_distribution }}" alt="Error Distribution">
                </div>
                {% endif %}
            </div>
        </body>
        </html>
        """
        
        # Generate HTML
        template = Template(template_str)
        html_content = template.render(**template_data)
        
        # Save HTML file with UTF-8 encoding
        output_path = os.path.join('reports', f'{model_name}_{data_type.lower()}_analysis.html')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return output_path

    def generate_general_html_report(self, correlation_matrix, correlation_plot, pca_results, pca_plots, target_name):
        """Generate general HTML report with correlation and PCA analysis"""
        # Create reports directory if it doesn't exist
        os.makedirs('reports', exist_ok=True)
        
        # Update plot paths to use the correct path (plots directory)
        updated_plots = {}
        # Update correlation plot path
        if correlation_plot:
            updated_plots['correlation'] = os.path.join('plots', os.path.basename(correlation_plot))
        
        # Update PCA plot paths
        updated_pca_plots = {}
        for plot_name, plot_path in pca_plots.items():
            if plot_path:
                updated_pca_plots[plot_name] = os.path.join('plots', os.path.basename(plot_path))
        
        # Calculate cumulative variance ratio
        cumulative_variance = np.cumsum(pca_results['explained_variance_ratio'])
        
        # Prepare template data with required Python functions
        template_data = {
            'target_name': target_name,
            'correlation_matrix': correlation_matrix,
            'correlation_plot': updated_plots.get('correlation', ''),
            'pca_results': {
                'explained_variance_ratio': pca_results['explained_variance_ratio'],
                'cumulative_variance_ratio': cumulative_variance,
                'components': pca_results['components'],
                'feature_names': [col for col in correlation_matrix.columns if col != target_name]
            },
            'pca_plots': updated_pca_plots,
            'zip': zip,
            'enumerate': enumerate,
            'range': range,
            'len': len,
            'format_float': lambda x: f"{x:.3f}"
        }
        
        # Load and customize the template
        template = Template("""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>General Analysis Report</title>
            <style>
                body {
                    font-family: Arial, sans-serif;
                    margin: 40px;
                    line-height: 1.6;
                }
                .section {
                    margin-bottom: 30px;
                    background: #fff;
                    padding: 20px;
                    border-radius: 8px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                }
                table {
                    border-collapse: collapse;
                    width: 100%;
                    margin: 20px 0;
                }
                th, td {
                    border: 1px solid #ddd;
                    padding: 12px;
                    text-align: center;
                }
                th {
                    background-color: #f2f2f2;
                }
                tr:nth-child(even) {
                    background-color: #f9f9f9;
                }
                img {
                    max-width: 100%;
                    height: auto;
                    margin: 20px 0;
                    border-radius: 4px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                }
                .plot-container {
                    text-align: center;
                    margin: 20px 0;
                }
                h1, h2, h3 {
                    color: #2c3e50;
                    border-bottom: 2px solid #eee;
                    padding-bottom: 10px;
                }
            </style>
        </head>
        <body>
            <h1>General Analysis Report</h1>
            
            <div class="section">
                <h2>Correlation Analysis</h2>
                <p>Correlation matrix for all variables including the target variable ({{ target_name }}):</p>
                <table>
                    <tr>
                        <th></th>
                        {% for col in correlation_matrix.columns %}
                        <th>{{ col }}</th>
                        {% endfor %}
                    </tr>
                    {% for index, row in correlation_matrix.iterrows() %}
                    <tr>
                        <th>{{ index }}</th>
                        {% for value in row %}
                        <td>{{ format_float(value) }}</td>
                        {% endfor %}
                    </tr>
                    {% endfor %}
                </table>
                
                {% if correlation_plot %}
                <div class="plot-container">
                    <h3>Correlation Heatmap</h3>
                    <img src="{{ correlation_plot }}" alt="Correlation Heatmap">
                </div>
                {% endif %}
            </div>
            
            <div class="section">
                <h2>Principal Component Analysis (PCA)</h2>
                
                <h3>Explained Variance Ratio</h3>
                <table>
                    <tr>
                        <th>Component</th>
                        <th>Explained Variance Ratio</th>
                        <th>Cumulative Explained Variance Ratio</th>
                    </tr>
                    {% for i in range(len(pca_results['explained_variance_ratio'])) %}
                    <tr>
                        <td>PC{{ i + 1 }}</td>
                        <td>{{ format_float(pca_results['explained_variance_ratio'][i]) }}</td>
                        <td>{{ format_float(pca_results['cumulative_variance_ratio'][i]) }}</td>
                    </tr>
                    {% endfor %}
                </table>
                
                <h3>Component Loadings</h3>
                <table>
                    <tr>
                        <th>Feature</th>
                        {% for i in range(pca_results['components'].shape[1]) %}
                        <th>PC{{ i + 1 }}</th>
                        {% endfor %}
                    </tr>
                    {% for feature, loadings in zip(pca_results['feature_names'], pca_results['components']) %}
                    <tr>
                        <td>{{ feature }}</td>
                        {% for loading in loadings %}
                        <td>{{ format_float(loading) }}</td>
                        {% endfor %}
                    </tr>
                    {% endfor %}
                </table>
                
                {% if pca_plots.get('variance') %}
                <div class="plot-container">
                    <h3>PCA Variance Plot</h3>
                    <img src="{{ pca_plots['variance'] }}" alt="PCA Variance Plot">
                </div>
                {% endif %}
                
                {% if pca_plots.get('loadings') %}
                <div class="plot-container">
                    <h3>PCA Loadings Plot</h3>
                    <img src="{{ pca_plots['loadings'] }}" alt="PCA Loadings Plot">
                </div>
                {% endif %}
            </div>
        </body>
        </html>
        """)
        
        # Generate HTML
        html_content = template.render(**template_data)
        
        # Save HTML file
        output_path = os.path.join('reports', 'general_analysis.html')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return output_path 